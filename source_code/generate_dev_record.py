# =============================================================
# generate_dev_record.py — FundusAI-Edu 开发记录 (DOCX)
# 输出: F:/demo/source_code/FundusAI-Edu_开发记录.docx
# =============================================================

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from pathlib import Path

OUT = Path(__file__).parent / "FundusAI-Edu_开发记录.docx"
IMG = Path(__file__).parent / "images"


def set_style(doc):
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    for i in range(1, 4):
        hs = doc.styles[f'Heading {i}']
        hs.font.color.rgb = RGBColor(0x0D, 0x11, 0x17)
        hs.font.bold = True
        hs.font.size = {1: Pt(22), 2: Pt(16), 3: Pt(13)}[i]


def h(doc, text, level=1):
    doc.add_heading(text, level=level)


def p(doc, text, bold=False, indent=False):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    if indent:
        para.paragraph_format.left_indent = Cm(1)
    return para


def bullet(doc, text, level=0):
    para = doc.add_paragraph(text, style='List Bullet')
    if level > 0:
        para.paragraph_format.left_indent = Cm(1 + level * 1.5)


def code(doc, text):
    for line in text.strip().split('\n'):
        para = doc.add_paragraph()
        run = para.add_run(line)
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x58, 0xA6, 0xFF)


def table(doc, headers, rows):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = 'Light Grid Accent 1'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, hd in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.text = hd
        for pp in cell.paragraphs:
            pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in pp.runs:
                run.bold = True
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            tbl.rows[ri + 1].cells[ci].text = str(val)
    doc.add_paragraph()
    return tbl


def img(doc, path, caption="", width=5.5):
    if path.exists():
        doc.add_picture(str(path), width=Inches(width))
        if caption:
            cp = doc.add_paragraph()
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cp.add_run(caption)
            run.italic = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x8B, 0x94, 0x9E)
        doc.add_paragraph()


def build():
    doc = Document()
    set_style(doc)

    # ====== 封面 ======
    for _ in range(6):
        doc.add_paragraph()
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("FundusAI-Edu"); r.bold = True; r.font.size = Pt(36); r.font.color.rgb = RGBColor(0x58, 0xA6, 0xFF)
    s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = s.add_run("项目开发记录"); r.font.size = Pt(20); r.font.color.rgb = RGBColor(0x8B, 0x94, 0x9E)
    doc.add_paragraph()
    for line in ["版本 v1.0", "开发周期: 2026年6月", "文档生成日期: 2026-06-07"]:
        pp = doc.add_paragraph(); pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = pp.add_run(line); run.font.size = Pt(12); run.font.color.rgb = RGBColor(0x8B, 0x94, 0x9E)
    doc.add_page_break()

    # ====== 目录 ======
    h(doc, "目录")
    for item in [
        "1. 项目概述", "2. 系统架构", "3. 技术选型", "4. 模块开发",
        "   4.1 眼底图像分析模块", "   4.2 血管拓扑分析模块",
        "   4.3 RAG 智能问答模块", "   4.4 科研选题生成模块",
        "   4.5 论文写作辅助模块", "   4.6 学习分析模块",
        "5. UI/UX 设计", "6. 数据处理管线", "7. 关键问题与解决方案",
        "8. 版本迭代", "9. 代码统计", "10. 待办事项",
    ]:
        pp = doc.add_paragraph(item); pp.paragraph_format.space_after = Pt(3)
    doc.add_page_break()

    # ====== 1. 项目概述 ======
    h(doc, "1. 项目概述")
    p(doc, "FundusAI-Edu 是一款面向医学教学与科研的眼底图像人工智能分析平台，集成深度学习图像处理和自然语言处理技术。项目目标：为学生和研究人员提供一站式的眼底图像分析工具链。")
    p(doc, "核心目标：")
    bullet(doc, "让医学生通过可视化交互理解眼底图像AI分析原理")
    bullet(doc, "降低科研门槛，AI辅助选题、实验设计和论文写作")
    bullet(doc, "提供可复用的教学演示环境，适配多种教学模式")

    table(doc, ["项目属性", "值"], [
        ["项目名称", "FundusAI-Edu"],
        ["开发语言", "Python 3.10+"],
        ["框架", "Streamlit + PyTorch + LangChain"],
        ["代码量", "~4000 行"],
        ["模块数", "6 个核心模块 + 入口 + 配置"],
        ["依赖包", "28 个"],
        ["目标用户", "高职高专 / 本科生 / 硕士 / 博士"],
    ])

    doc.add_page_break()

    # ====== 2. 系统架构 ======
    h(doc, "2. 系统架构")
    p(doc, "平台采用四层架构设计，从用户交互到底层数据存储逐层解耦。")
    img(doc, IMG / "arch_diagram.png", "图1: FundusAI-Edu 四层系统架构图")

    p(doc, "架构分层说明：")
    table(doc, ["层级", "组件", "技术"], [
        ["展示层", "Streamlit Web UI + pywebview 桌面窗口", "St.markdown + Custom CSS"],
        ["业务层", "6 个功能模块", "OpenCV / Frangi / PyTorch U-Net"],
        ["AI 层", "DeepSeek LLM + ChromaDB + PyTorch", "LangChain RAG"],
        ["数据层", "SQLite + 知识库文件 + 眼底图像", "chromadb / PIL / pypdf"],
    ])

    img(doc, IMG / "deps.png", "图2: 模块依赖关系图")
    p(doc, "app.py 作为主入口依赖全部 6 个功能模块，config.py 作为全局配置中心被所有模块依赖。模块之间保持独立，无交叉引用。")

    doc.add_page_break()

    # ====== 3. 技术选型 ======
    h(doc, "3. 技术选型")
    table(doc, ["分类", "技术", "选型理由"], [
        ["UI 框架", "Streamlit 1.58", "纯 Python 构建 Web 应用，无需前端代码"],
        ["桌面", "pywebview 6.2", "嵌入 Streamlit 到原生窗口，体验接近桌面应用"],
        ["打包", "PyInstaller 6", "一键打包为 Windows .exe 文件"],
        ["深度学习", "PyTorch (CPU)", "U-Net 血管分割，CPU 模式兼容性最佳"],
        ["图像处理", "OpenCV + scikit-image", "行业标准，Frangi 滤波 / CLAHE / 形态学"],
        ["拓扑分析", "NetworkX + skimage", "图论骨架化 + 六维特征提取"],
        ["RAG", "LangChain + ChromaDB", "文档检索增强生成，离线知识库"],
        ["LLM", "DeepSeek (OpenAI SDK)", "国产大模型，性价比高，兼容 OpenAI API"],
        ["嵌入", "sentence-transformers", "文本向量化，知识库检索"],
        ["可视化", "matplotlib + plotly + seaborn", "雷达图 / 骨架图 / 统计图"],
        ["数据库", "SQLite", "轻量级，零配置，适合教学场景"],
        ["文档解析", "pypdf + python-docx", "知识库文档上传解析"],
    ])

    doc.add_page_break()

    # ====== 4. 模块开发 ======
    h(doc, "4. 模块开发")

    # 4.1
    h(doc, "4.1 眼底图像分析模块 (fundus_analysis.py)", 2)
    p(doc, "核心功能：血管分割 + 病变识别 + 可视化叠加。")
    p(doc, "技术要点：")
    bullet(doc, "U-Net 模型定义：DoubleConv → Down → Up → OutConv 标准结构（3通道输入，1通道输出）")
    bullet(doc, "演示模式：多尺度黑帽变换(7/15/31) + Frangi 滤波 + 自适应百分位阈值")
    bullet(doc, "病变识别：HSV 色彩空间规则 — 出血点(红) / 硬性渗出(黄白) / DR 四级分级")
    bullet(doc, "血管叠加：二值化掩码 > 绿色通道 + cv2.addWeighted 透明叠加")
    bullet(doc, "图像预处理：CLAHE 增强绿通道对比度 → 归一化 → PyTorch Tensor")
    p(doc, "关键代码片段：")
    code(doc, """class FundusAnalyzer:
    def _demo_segment(self, image):
        gray = cv2.cvtColor(image, cv2.COLOR_RGB2GRAY)
        enhanced = clahe.apply(gray)
        blackhat = cv2.morphologyEx(enhanced, cv2.MORPH_BLACKHAT, kernel)
        frangi = filters.frangi(1.0 - enhanced/255.0, sigmas=range(1,5))
        combined = 0.6 * blackhat/255.0 + 0.4 * frangi_norm
        binary = (combined > threshold).astype(np.uint8) * 255""")

    # 4.2
    h(doc, "4.2 血管拓扑分析模块 (topology_analysis.py)", 2)
    p(doc, "核心功能：骨架提取 → 拓扑图构建 → 六维特征计算 + 雷达图/骨架三连图。")
    bullet(doc, "骨架化：Lee 算法 (skimage.morphology.skeletonize) + 小连通域过滤(<50px)")
    bullet(doc, "拓扑图：8-邻域构建 NetworkX 无向图，节点=骨架像素，边=相邻骨架")
    bullet(doc, "六维特征：密度 | 分叉点 | 终末点 | 平均宽度(距离变换) | 迂曲度(路径/直线) | 分形维数(盒计数)")
    bullet(doc, "骨架可视化：三连图 — 原图 / 绿色血管标注 / 红色骨架叠加(dilated disk(1))")
    bullet(doc, "雷达图：matplotlib polar 直角坐标，六轴归一化，品牌蓝色填充")

    # 4.3
    h(doc, "4.3 RAG 智能问答模块 (rag_chat.py)", 2)
    p(doc, "核心功能：文档检索增强生成 + DeepSeek 大模型多轮对话。")
    bullet(doc, "ChromaDB 向量库存储知识文档，sentence-transformers 做文本嵌入")
    bullet(doc, "检索策略：相似度 Top-K → 拼接上下文 → LLM 润色回答")
    bullet(doc, "离线兜底：无 API Key 时用关键词匹配内置知识库，功能不中断")
    bullet(doc, "知识库管理：支持 PDF/TXT/DOCX 上传，自动分块(chunk_size=500) 入库")

    doc.add_page_break()

    # 4.4
    h(doc, "4.4 科研选题生成模块 (topic_generator.py)", 2)
    p(doc, "核心功能：AI 自动生成选题 → 研究方案 → 实验报告。")
    bullet(doc, "输入：研究方向 + 学术层次 + 自定义需求")
    bullet(doc, "输出：选题列表(含创新点/数据集建议) + 详细研究方案 + 实验设计模板")
    bullet(doc, "学术层次适配：高职高专 → 应用型；本科 → 综述+实验；硕士 → 算法改进；博士 → 理论创新")

    # 4.5
    h(doc, "4.5 论文写作辅助模块 (paper_assistant.py)", 2)
    p(doc, "支持 9 种论文章节类型自动生成，中英双语。")
    table(doc, ["章节", "生成内容", "语言"], [
        ["Abstract", "研究背景 + 方法 + 结果 + 结论", "中/英"],
        ["Introduction", "领域背景 + 问题陈述 + 贡献", "中/英"],
        ["Related Work", "文献综述 + 对比分析", "中/英"],
        ["Method", "算法描述 + 流程图", "中/英"],
        ["Experiments", "数据集 + 评测指标 + 结果分析", "中/英"],
        ["Discussion", "结果解释 + 局限 + 展望", "中/英"],
        ["Conclusion", "总结 + 贡献 + 未来工作", "中/英"],
        ["References", "10-20 篇推荐文献", "英文"],
        ["Journals", "推荐投稿期刊 + 影响因子", "英文"],
    ])

    # 4.6
    h(doc, "4.6 学习分析模块 (learning_analytics.py)", 2)
    p(doc, "SQLite 自动记录用户操作，生成个人学习画像。")
    bullet(doc, "数据库表：users, actions, analyses, chat_sessions, topic_sessions")
    bullet(doc, "统计维度：操作次数 / 分析次数 / 模块使用频率 / 学习时长 / 进步曲线")
    bullet(doc, "可视化：柱状图(plotly) + 折线图(学习趋势) + 饼图(模块分布)")

    doc.add_page_break()

    # ====== 5. UI/UX 设计 ======
    h(doc, "5. UI/UX 设计")
    img(doc, IMG / "ui_nav.png", "图3: UI 页面导航结构")
    p(doc, "设计理念：暗色主题(0d1117) + 品牌蓝(58a6ff) + 高对比度侧边栏。")
    bullet(doc, "侧边栏：深色渐变背景(0d1117→161b22)，CSS 覆盖 Streamlit 默认白色输入组件")
    bullet(doc, "所有输入组件：独立深色底(1c2333) + 2px 边框(4a5568) + 焦点蓝光(58a6ff)")
    bullet(doc, "英雄标题：品牌蓝(58a6ff) + text-shadow 发光 + 1.5 行高避免截断")
    bullet(doc, "响应式：wide layout + columns 布局，14 寸至 27 寸自适应")
    p(doc, "关键 CSS 规则数量：~160 行自定义样式，覆盖了 Streamlit 全部内置组件。")

    # ====== 6. 数据处理管线 ======
    h(doc, "6. 数据处理管线")
    img(doc, IMG / "pipeline.png", "图4: 血管分割算法管线")
    p(doc, "管线分为 5 个阶段：")
    table(doc, ["阶段", "操作", "输出"], [
        ["1. 输入", "PIL.Image.open → RGB → Gray", "灰度图"],
        ["2. 增强", "CLAHE(clipLimit=2.0) → 反转", "对比度增强图"],
        ["3. 滤波", "Multi-scale BlackHat(7/15/31) + Frangi", "响应图融合"],
        ["4. 阈值", "百分位自适应 → 形态学 Open/Close", "二值掩码"],
        ["5. 后处理", "骨架化 + 距离变换 + 图论分析", "六维特征"],
    ])

    doc.add_page_break()

    # ====== 7. 关键问题与解决方案 ======
    h(doc, "7. 关键问题与解决方案")
    table(doc, ["问题", "根因", "解决方案"], [
        ["图片上传显示红色感叹号", "st.file_uploader type参数对大小写敏感(.JPG被拒)", "取消type限制，手动用.lower()校验后缀"],
        ["Connection lost 连接断开", "沙盒拒绝写入~/.streamlit/machine_id_v4", "设置STREAMLIT_HOME重定向到项目目录"],
        ["Connection lost 反复出现", "端口8501有10+僵尸进程残留", "PowerShell Get-NetTCPConnection逐pid kill"],
        ["PyInstaller打包报GBK编码", "代码中含Unicode emoji(\u2705等)", "全部替换为ASCII [OK]/[WARN]"],
        ["血管分割掩码全黑/全白", "Frangi参数设反(black_ridges=False)", "图像反转+black_ridges=False + 多尺度黑帽融合"],
        ["骨架可视化空白", "掩码尺寸(512x512)与原图不一致", "统一缩放到原图尺寸再绘图"],
        ["标题文字被截断", "-webkit-background-clip渐变渲染bug", "改用实色#58a6ff + overflow:visible"],
        ["侧边栏白底组件对比度低", "Streamlit默认白底在暗色侧边栏不兼容", "~160行CSS覆盖所有组件样式"],
    ])

    # ====== 8. 版本迭代 ======
    h(doc, "8. 版本迭代")
    table(doc, ["版本", "日期", "变更"], [
        ["v0.1", "2026-06-07", "初始化项目，生成6个模块 + app.py + config.py"],
        ["v0.2", "2026-06-07", "暗色主题CSS + 侧边栏全局高对比度"],
        ["v0.3", "2026-06-07", "图片上传修复 + 清除按钮 + 异常保护"],
        ["v0.4", "2026-06-07", "Connection lost 修复（端口清理 + 沙盒权限）"],
        ["v0.5", "2026-06-07", "血管分割算法重写（多尺度黑帽+Frangi融合）"],
        ["v0.6", "2026-06-07", "骨架可视化重写（原图叠加标注）"],
        ["v1.0", "2026-06-07", "桌面弹窗模式 + PyInstaller 打包 + 使用手册"],
    ])

    # ====== 9. 代码统计 ======
    h(doc, "9. 代码统计")
    table(doc, ["文件", "行数", "职责"], [
        ["app.py", "1365", "Streamlit 主入口 + 6页面 + CSS"],
        ["config.py", "126", "全局配置中心"],
        ["fundus_analysis.py", "375", "U-Net + Frangi 血管分割"],
        ["topology_analysis.py", "389", "骨架化 + 六维特征"],
        ["rag_chat.py", "342", "ChromaDB + RAG 问答"],
        ["topic_generator.py", "314", "科研选题 + 方案生成"],
        ["paper_assistant.py", "286", "论文框架 + 章节辅助"],
        ["learning_analytics.py", "466", "SQLite + 学习画像"],
        ["run.py / launcher.py", "230", "桌面弹窗启动器"],
        ["build.py + spec", "130", "PyInstaller 打包"],
        ["requirements.txt", "58", "28 个依赖包"],
        ["合计", "~4080", ""],
    ])

    doc.add_page_break()

    # ====== 10. 待办事项 ======
    h(doc, "10. 待办事项")
    table(doc, ["优先级", "事项", "说明"], [
        ["P0", "U-Net 权重训练", "使用 DRIVE/STARE 数据集训练，当前仅演示模式"],
        ["P0", "图片上传稳定性", "偶发 WebSocket 断连，需进一步调查"],
        ["P1", "多语言支持", "增加英文版 UI"],
        ["P1", "用户认证", "添加登录/注册 + 数据隔离"],
        ["P2", "云端部署", "部署到 Streamlit Cloud / HuggingFace Spaces"],
        ["P2", "移动端适配", "响应式设计适配手机和平板"],
        ["P3", "插件系统", "支持第三方模块扩展"],
    ])

    # 尾页
    doc.add_page_break()
    h(doc, "开发团队")
    p(doc, "项目：FundusAI-Edu 眼底图像AI教学与科研辅助平台")
    p(doc, "版本：v1.0")
    p(doc, "Git 提交记录：F:\\demo\\.git")
    p(doc, "© 2026 FundusAI-Edu 团队")

    doc.save(str(OUT))
    print(f"[OK] 开发记录已生成: {OUT}")


if __name__ == "__main__":
    build()
