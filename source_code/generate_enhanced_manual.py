# =============================================================
# generate_enhanced_manual.py — 生成增强版用户使用手册
# 包含 UI 截图 + 详细功能说明 + 快速入门指南
# =============================================================

import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── 输出路径 ──
OUTPUT = os.path.join(os.path.dirname(__file__), "FundusAI-Edu_用户使用手册_v2.docx")

# ── 截图路径 ──
SCREENSHOT_DIR = os.path.join(os.path.dirname(__file__), "..", "slides", "imgs")
CLIPBOARD_IMG = os.path.join(os.path.expanduser("~"), ".workbuddy", "clipboard-images",
                              "clipboard-2026-06-07T14-42-23-358Z-74c2fa26.png")

SCREENSHOTS = {
    "main_page":     os.path.join(SCREENSHOT_DIR, "ui_01_upload.png"),
    "lesion_result": os.path.join(SCREENSHOT_DIR, "ui_02_lesion.png"),
    "topo_radar":    os.path.join(SCREENSHOT_DIR, "ui_03_topo1.png"),
    "topo_features": os.path.join(SCREENSHOT_DIR, "ui_04_topo2.png"),
    "api_config":    os.path.join(SCREENSHOT_DIR, "ui_05_api.png"),
    "learning":      os.path.join(SCREENSHOT_DIR, "ui_06_learning.png"),
    "topic_gen":     os.path.join(SCREENSHOT_DIR, "ui_07_topic.png"),
    "current_page":  CLIPBOARD_IMG,
}

# ── 辅助函数 ──

def set_cell_shading(cell, color_hex):
    """设置单元格背景色"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def add_colored_heading(doc, text, level=1, color=None):
    """添加带颜色的标题"""
    heading = doc.add_heading(text, level=level)
    if color:
        for run in heading.runs:
            run.font.color.rgb = RGBColor(*color)
    return heading


def add_styled_paragraph(doc, text, style='Normal', bold=False, italic=False,
                          font_size=None, color=None, alignment=None, space_after=None):
    """添加带样式的段落"""
    para = doc.add_paragraph(style=style)
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    if font_size:
        run.font.size = Pt(font_size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    if alignment is not None:
        para.alignment = alignment
    if space_after is not None:
        para.paragraph_format.space_after = Pt(space_after)
    return para


def add_image_safe(doc, image_path, width_inches=5.5, caption=None):
    """安全添加图片（检查文件是否存在）"""
    if image_path and os.path.exists(image_path):
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        run.add_picture(image_path, width=Inches(width_inches))

        if caption:
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap_run = cap.add_run(caption)
            cap_run.font.size = Pt(9)
            cap_run.font.color.rgb = RGBColor(100, 100, 100)
            cap_run.italic = True
        return True
    else:
        doc.add_paragraph(f'[ 截图未找到: {os.path.basename(image_path) if image_path else "N/A"} ]')
        return False


def add_info_box(doc, text, bg_color="E8F4FD"):
    """添加信息提示框"""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)
    # 用表格模拟提示框
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, bg_color)
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(f"💡 {text}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0, 51, 102)
    doc.add_paragraph()  # spacing


def add_step(doc, number, title, description):
    """添加操作步骤"""
    para = doc.add_paragraph()
    run_num = para.add_run(f"步骤 {number}：")
    run_num.bold = True
    run_num.font.size = Pt(11)
    run_title = para.add_run(title)
    run_title.bold = True
    run_title.font.size = Pt(11)

    desc = doc.add_paragraph(description)
    desc.paragraph_format.left_indent = Cm(1)
    for run in desc.runs:
        run.font.size = Pt(10)


def add_feature_card(doc, icon, title, description, screenshot_path=None, caption=None):
    """添加功能模块卡片（标题 + 描述 + 截图）"""
    # 标题
    h = doc.add_heading(f"{icon}  {title}", level=2)
    for run in h.runs:
        run.font.size = Pt(16)

    # 描述
    if description:
        p = doc.add_paragraph(description)
        p.paragraph_format.space_after = Pt(8)

    # 截图
    if screenshot_path:
        add_image_safe(doc, screenshot_path, width_inches=5.2, caption=caption)

    doc.add_paragraph()  # 分割


# =============================================================
# 创建文档
# =============================================================

doc = Document()

# ── 页面设置 ──
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

# ── 样式 ──
style = doc.styles['Normal']
font = style.font
font.name = 'Microsoft YaHei'
font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

for level in range(1, 4):
    h_style = doc.styles[f'Heading {level}']
    h_font = h_style.font
    h_font.name = 'Microsoft YaHei'
    h_style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    if level == 1:
        h_font.size = Pt(22)
        h_font.color.rgb = RGBColor(0, 51, 102)
    elif level == 2:
        h_font.size = Pt(16)
        h_font.color.rgb = RGBColor(0, 70, 127)
    elif level == 3:
        h_font.size = Pt(13)
        h_font.color.rgb = RGBColor(51, 102, 153)

# =============================================================
# 封面
# =============================================================

for _ in range(6):
    doc.add_paragraph()

title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_para.add_run("FundusAI-Edu")
title_run.font.size = Pt(36)
title_run.bold = True
title_run.font.color.rgb = RGBColor(0, 51, 102)

subtitle_para = doc.add_paragraph()
subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_run = subtitle_para.add_run("眼底图像AI教学与科研辅助平台")
sub_run.font.size = Pt(20)
sub_run.font.color.rgb = RGBColor(88, 166, 255)

doc.add_paragraph()

manual_para = doc.add_paragraph()
manual_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
manual_run = manual_para.add_run("用户使用手册 v2.0")
manual_run.font.size = Pt(16)
manual_run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()

date_para = doc.add_paragraph()
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_run = date_para.add_run("2026年6月 · 增强版（含UI截图）")
date_run.font.size = Pt(11)
date_run.font.color.rgb = RGBColor(140, 140, 140)

doc.add_page_break()

# =============================================================
# 目录页
# =============================================================

doc.add_heading("目  录", level=1)
toc_items = [
    ("1", "平台概述", "了解 FundusAI-Edu 的功能与定位"),
    ("2", "快速开始", "30 秒上手指南"),
    ("3", "眼底图像智能分析", "上传图像 → 血管分割 → 病变识别"),
    ("4", "血管拓扑特征分析", "六维特征提取与临床解读"),
    ("5", "AI科研导师", "RAG 驱动的医学知识问答"),
    ("6", "科研选题生成器", "AI 智能推荐研究方向"),
    ("7", "论文写作辅助", "框架生成 + 章节写作 + 润色"),
    ("8", "学习分析", "行为追踪与个人学习画像"),
    ("9", "API 配置", "接入大语言模型服务"),
    ("10", "常见问题", "FAQ 与故障排除"),
]
for num, title, desc in toc_items:
    p = doc.add_paragraph()
    run_num = p.add_run(f"{num}. ")
    run_num.bold = True
    run_num.font.size = Pt(11)
    run_title = p.add_run(f"{title}")
    run_title.font.size = Pt(11)
    run_desc = p.add_run(f"  — {desc}")
    run_desc.font.size = Pt(9)
    run_desc.font.color.rgb = RGBColor(140, 140, 140)

doc.add_page_break()

# =============================================================
# 第1章：平台概述
# =============================================================

add_colored_heading(doc, "1. 平台概述", level=1, color=(0, 51, 102))

add_styled_paragraph(doc,
    "FundusAI-Edu 是一款面向医学教学与科研的眼底图像 AI 分析平台。"
    "平台集成深度学习图像分割、图论拓扑分析、大语言模型（LLM）问答、"
    "以及学习行为分析等能力，为眼科、医学影像和生物医学工程等专业"
    "的师生提供一站式智能辅助工具。",
    font_size=10.5, space_after=6)

add_styled_paragraph(doc,
    "适用对象：高职高专 / 本科生 / 硕士研究生 / 博士研究生",
    font_size=10.5, color=(0, 70, 127), bold=True, space_after=12)

# 平台截图
add_image_safe(doc, SCREENSHOTS["main_page"], width_inches=5.5,
               caption="图1-1  FundusAI-Edu 平台主界面（眼底图像分析页）")

doc.add_paragraph()

# 核心功能列表
add_colored_heading(doc, "核心功能", level=2)

features = [
    ("🩸 血管自动分割", "基于 U-Net 深度学习模型，自动化提取视网膜血管网络，支持可视化叠加显示"),
    ("🔎 病变智能识别", "检测出血点、硬性渗出等 DR（糖尿病视网膜病变）特征，提供分级评估"),
    ("🌐 拓扑特征分析", "计算血管密度、分叉点、终末点、平均宽度、迂曲度、分形维数六项量化指标"),
    ("🤖 AI 知识问答", "基于 RAG（检索增强生成）技术，结合医学文献知识库提供专业问答"),
    ("🎯 科研选题推荐", "根据用户研究方向与兴趣，AI 智能生成个性化研究课题与研究方案"),
    ("📝 论文写作辅助", "提供论文框架生成、逐章节写作指导和学术语言润色服务"),
    ("📊 学习行为分析", "自动追踪用户平台操作，生成多维度学习画像和能力雷达图"),
]

for icon_title, desc in features:
    p = doc.add_paragraph()
    run = p.add_run(f"{icon_title}  ")
    run.bold = True
    run.font.size = Pt(10.5)
    run2 = p.add_run(desc)
    run2.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(4)

# 技术架构
doc.add_paragraph()
add_colored_heading(doc, "技术栈", level=2)

techs = [
    ("🐍 Python", "核心开发语言"),
    ("⚡ Streamlit", "Web 交互界面"),
    ("🔥 PyTorch", "深度学习框架（U-Net）"),
    ("🧠 DeepSeek", "大语言模型（LLM）"),
    ("🗄️ ChromaDB", "向量数据库（RAG 知识库）"),
    ("📐 NetworkX", "图论拓扑分析"),
]

table = doc.add_table(rows=1, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = "技术组件"
hdr[1].text = "用途说明"
for cell in hdr:
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(10)

for tech_name, tech_desc in techs:
    row = table.add_row()
    row.cells[0].text = tech_name
    row.cells[1].text = tech_desc
    for cell in row.cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.size = Pt(10)

doc.add_page_break()

# =============================================================
# 第2章：快速开始
# =============================================================

add_colored_heading(doc, "2. 快速开始", level=1, color=(0, 51, 102))

add_info_box(doc, "首次使用？跟随以下 4 步，30 秒即可上手！")

add_step(doc, 1, "访问平台",
         "在浏览器中打开 http://localhost:8501 ，进入 FundusAI-Edu 主页。\n"
         "主页展示了 6 大功能模块卡片和技术架构概览。")
add_step(doc, 2, "侧边栏配置",
         "在左侧边栏中填写个人信息：用户名、学历层次、感兴趣的研究方向。\n"
         "这些信息将用于个性化学习和学习画像生成。")
add_step(doc, 3, "选择功能模块",
         "在侧边栏的「功能导航」中选择需要使用的模块：\n"
         "眼底图像分析 | 血管拓扑分析 | AI科研导师 | 科研选题生成器 | 论文写作辅助 | 学习分析")
add_step(doc, 4, "开始分析",
         "以「眼底图像分析」为例：上传一张 JPG/PNG 格式的眼底照片，\n"
         "点击「开始分析」按钮，等待 AI 完成血管分割、病变识别和拓扑特征提取。")

doc.add_paragraph()
add_image_safe(doc, SCREENSHOTS["current_page"], width_inches=5.5,
               caption="图2-1  快速入门：左侧上传图像，右侧查看分析结果")

doc.add_page_break()

# =============================================================
# 第3章：眼底图像智能分析
# =============================================================

add_colored_heading(doc, "3. 眼底图像智能分析", level=1, color=(0, 51, 102))

add_styled_paragraph(doc,
    "眼底图像智能分析是平台的核心功能，支持用户上传彩色眼底照片，"
    "通过 AI 模型自动完成血管分割、病变识别和拓扑特征提取，"
    "一站式获得全面的眼底分析报告。",
    font_size=10.5, space_after=10)

add_colored_heading(doc, "3.1 操作流程", level=2)

add_step(doc, 1, "上传图像",
         "在左侧区域点击「Browse files」或拖拽图像到上传区域。\n"
         "支持格式：JPG、PNG、BMP、TIFF、WebP\n"
         "建议使用高质量的彩色眼底照片以获得最佳分析效果。")
add_step(doc, 2, "预览与确认",
         "上传成功后，系统会显示原始图像缩略图，\n"
         "同时展示图像尺寸（像素）和文件大小（MB）信息。")
add_step(doc, 3, "开始分析",
         "点击蓝色「🚀 开始分析」按钮启动 AI 分析流水线：\n"
         "① 血管分割 (20%) → ② 病变识别 (50%) → ③ 拓扑分析 (70%) → ④ 可视化 (100%)\n"
         "整个过程约需 5-15 秒，取决于图像大小和 CPU 性能。")
add_step(doc, 4, "查看结果",
         "分析完成后，右侧出现三个结果标签页：\n"
         "「🔍 血管分割」「📈 病变识别」「🌐 拓扑特征」")

add_image_safe(doc, SCREENSHOTS["main_page"], width_inches=5.2,
               caption="图3-1  图像上传界面与功能说明")

doc.add_paragraph()

add_colored_heading(doc, "3.2 血管分割", level=2)
add_styled_paragraph(doc,
    "系统使用 U-Net 深度学习模型完成血管分割，并通过多尺度血管增强算法"
    "（Multi-Scale BlackHat + Frangi Filter）提高分割精度。结果以三个视图展示：",
    font_size=10, space_after=6)

views = [
    "血管掩码（绿色 = 血管，黑色 = 背景）：直观展示分割结果",
    "血管叠加图（半透明绿色叠加在原始图像上）：便于对比检视",
    "骨架可视化（提取血管中心线）：为拓扑分析提供基础",
]
for v in views:
    p = doc.add_paragraph(f"  • {v}", style='List Bullet')
    for run in p.runs:
        run.font.size = Pt(10)

add_colored_heading(doc, "3.3 病变识别", level=2)
add_styled_paragraph(doc,
    "系统基于 HSV 色彩空间规则检测出血点和硬性渗出，自动评估糖尿病视网膜病变"
    "（DR）分级和出血/渗出风险。",
    font_size=10, space_after=6)

add_image_safe(doc, SCREENSHOTS["lesion_result"], width_inches=5.2,
               caption="图3-2  病变识别结果：DR 分级 + 风险评估")

doc.add_paragraph()
add_styled_paragraph(doc,
    "DR 分级标准：No DR（无病变）、Mild DR（轻度）、Moderate DR（中度）、Severe DR（重度）",
    font_size=9, color=(100, 100, 100), italic=True)

add_colored_heading(doc, "3.4 拓扑特征", level=2)
add_styled_paragraph(doc,
    "系统自动计算六项血管拓扑特征：血管密度、分叉点数量、终末点数量、"
    "平均血管宽度、血管迂曲度、分形维数。每项特征均附有临床参考范围。",
    font_size=10, space_after=6)

doc.add_page_break()

# =============================================================
# 第4章：血管拓扑特征分析
# =============================================================

add_colored_heading(doc, "4. 血管拓扑特征分析", level=1, color=(0, 51, 102))

add_styled_paragraph(doc,
    "血管拓扑特征分析模块提供六维特征的详细展示和临床意义解读。"
    "该模块需要先完成眼底图像分析（或导入已有分析结果）。",
    font_size=10.5, space_after=10)

add_image_safe(doc, SCREENSHOTS["topo_features"], width_inches=5.2,
               caption="图4-1  六大拓扑特征指标卡片")

doc.add_paragraph()

add_colored_heading(doc, "4.1 六维特征说明", level=2)

topo_features = [
    ("血管密度", "vessel_density", "0.08 ~ 0.18", "血管面积与总面积的比值，反映视网膜供血状况"),
    ("分叉点数量", "bifurcation_count", "随图像变化", "血管网络的分支节点数，反映血管网络复杂度"),
    ("终末点数量", "endpoint_count", "随图像变化", "血管末端节点数，异常增多提示新生血管可能"),
    ("平均血管宽度", "avg_vessel_width", "3 ~ 8 px", "反映血管口径，变窄可能与高血压等疾病相关"),
    ("血管迂曲度", "tortuosity", "1.0 ~ 1.3", "血管实际长度与直线距离的比值，升高提示血管形态异常"),
    ("分形维数", "fractal_dimension", "1.5 ~ 1.7", "血管网络空间填充能力，可用于评估视网膜微循环健康状态"),
]

table = doc.add_table(rows=1, cols=4)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = "特征名称"
hdr[1].text = "正常范围"
hdr[2].text = "临床意义"
hdr[3].text = "参数名"
for cell in hdr:
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(9)

for name, param, normal, clinical in topo_features:
    row = table.add_row()
    row.cells[0].text = name
    row.cells[1].text = normal
    row.cells[2].text = clinical
    row.cells[3].text = param
    for cell in row.cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.size = Pt(8.5)

doc.add_paragraph()

add_image_safe(doc, SCREENSHOTS["topo_radar"], width_inches=5.2,
               caption="图4-2  特征雷达图 + 临床意义解读")

doc.add_page_break()

# =============================================================
# 第5章：AI科研导师
# =============================================================

add_colored_heading(doc, "5. AI科研导师", level=1, color=(0, 51, 102))

add_styled_paragraph(doc,
    "AI科研导师模块基于 RAG（检索增强生成）技术，将专业医学文献知识库与大语言模型结合，"
    "为用户提供精准、有据可查的医学知识问答服务。",
    font_size=10.5, space_after=10)

add_colored_heading(doc, "5.1 功能特点", level=2)
features_rag = [
    "多轮对话：支持连续提问，保持对话上下文",
    "文献引用：回答附有知识来源引用，可追溯",
    "医学专精：知识库涵盖眼科、糖尿病视网膜病变、血管生物学等方向",
    "双语支持：支持中文和英文问答",
]
for f in features_rag:
    p = doc.add_paragraph(f"  • {f}", style='List Bullet')

add_colored_heading(doc, "5.2 使用方式", level=2)
add_step(doc, 1, "配置 API",
         "在左侧边栏中配置 DeepSeek API Key 和服务地址。\n"
         "如尚未配置，请参考第 9 章「API 配置」。")
add_step(doc, 2, "输入问题",
         "在聊天输入框中输入您关心的医学或科研问题，\n"
         "例如：「糖尿病视网膜病变的病理机制是什么？」")
add_step(doc, 3, "查看回答",
         "系统自动检索知识库中的相关文献片段，结合 LLM 生成回答。\n"
         "回答下方会显示引用的知识来源。")

doc.add_page_break()

# =============================================================
# 第6章：科研选题生成器
# =============================================================

add_colored_heading(doc, "6. 科研选题生成器", level=1, color=(0, 51, 102))

add_styled_paragraph(doc,
    "科研选题生成器根据用户的学历层次和研究方向，"
    "利用 AI 大语言模型智能生成个性化科研选题、研究方法和实验方案建议。",
    font_size=10.5, space_after=10)

add_image_safe(doc, SCREENSHOTS["topic_gen"], width_inches=5.2,
               caption="图6-1  科研选题生成器界面：配置 + 生成 + 研究方案")

doc.add_paragraph()

add_colored_heading(doc, "6.1 使用方式", level=2)
add_step(doc, 1, "配置研究偏好",
         "在左侧面板中选择学历层次（高职高专/本科/硕士/博士）\n"
         "和感兴趣的研究方向（DR 分级、血管分割、多模态融合等）。")
add_step(doc, 2, "调整参数",
         "可调整选题数量和推荐温度（创造性程度）。\n"
         "温度越高，选题越具创新性；温度越低，选题越务实。")
add_step(doc, 3, "生成选题",
         "点击「生成选题」按钮，AI 将生成包含标题、背景、创新点、\n"
         "研究方法和预期成果的完整研究方案。")

doc.add_page_break()

# =============================================================
# 第7章：论文写作辅助
# =============================================================

add_colored_heading(doc, "7. 论文写作辅助", level=1, color=(0, 51, 102))

add_styled_paragraph(doc,
    "论文写作辅助模块为科研论文写作提供全流程 AI 辅助，"
    "从框架搭建到逐章节写作，再到语言润色和期刊推荐。",
    font_size=10.5, space_after=10)

add_colored_heading(doc, "7.1 功能列表", level=2)
paper_features = [
    ("📋 框架生成", "根据论文标题和关键词，自动生成 IMRaD 标准结构框架"),
    ("✍️ 章节写作", "按 Introduction/Methods/Results/Discussion 逐章节生成初稿"),
    ("✨ 语言润色", "优化学术语言表达，提升论文可读性"),
    ("📰 期刊推荐", "根据论文主题推荐匹配的 SCI 期刊"),
]
for icon, desc in paper_features:
    p = doc.add_paragraph()
    run = p.add_run(f"{icon}  ")
    run.bold = True
    run2 = p.add_run(desc)
    run2.font.size = Pt(10)

add_colored_heading(doc, "7.2 使用流程", level=2)
add_step(doc, 1, "输入论文信息",
         "填写论文标题、关键词和研究领域。\n"
         "建议提供清晰、具体的标题以获得更精准的辅助。")
add_step(doc, 2, "选择辅助类型",
         "从下拉菜单选择需要的辅助类型：\n"
         "框架生成 / 章节写作 / 语言润色 / 期刊推荐")
add_step(doc, 3, "获取结果",
         "AI 生成对应的论文内容或建议，可直接复制使用。\n"
         "所有生成内容仅供参考，建议结合专业知识进行审阅和修改。")

doc.add_page_break()

# =============================================================
# 第8章：学习分析
# =============================================================

add_colored_heading(doc, "8. 学习分析", level=1, color=(0, 51, 102))

add_styled_paragraph(doc,
    "学习分析模块自动追踪用户在平台上的所有操作，"
    "生成多维度学习行为统计、可视化图表和个人学习画像。",
    font_size=10.5, space_after=10)

add_image_safe(doc, SCREENSHOTS["learning"], width_inches=5.2,
               caption="图8-1  学习分析仪表盘：统计数据 + 活动图表 + 学习画像")

doc.add_paragraph()

add_colored_heading(doc, "8.1 分析维度", level=2)
analytics_dims = [
    ("📈 操作统计", "分析次数、对话轮数、选题生成次数等"),
    ("📊 活跃度图表", "按日期统计平台使用频率（柱状图 + 趋势线）"),
    ("🍩 功能分布", "饼图展示各模块使用占比"),
    ("🎯 学习画像", "六维雷达图：选题、查证、分析、写作、综合、创新"),
]
for icon, desc in analytics_dims:
    p = doc.add_paragraph()
    run = p.add_run(f"{icon}：")
    run.bold = True
    run.font.size = Pt(10)
    run2 = p.add_run(desc)
    run2.font.size = Pt(10)

add_colored_heading(doc, "8.2 数据导出", level=2)
add_styled_paragraph(doc,
    "点击「⬇️ 下载CSV」按钮导出详细学习报告（CSV格式）；\n"
    "点击「📄 导出学习画像」按钮导出个人学习画像（JSON格式）。",
    font_size=10, space_after=6)

doc.add_page_break()

# =============================================================
# 第9章：API 配置
# =============================================================

add_colored_heading(doc, "9. API 配置", level=1, color=(0, 51, 102))

add_styled_paragraph(doc,
    "平台的部分 AI 功能（RAG 问答、选题生成、论文辅助）需要接入大语言模型服务。"
    "建议使用 DeepSeek API，以获得最佳兼容性和性价比。",
    font_size=10.5, space_after=10)

add_image_safe(doc, SCREENSHOTS["api_config"], width_inches=5.2,
               caption="图9-1  API 配置面板（侧边栏底部）")

doc.add_paragraph()

add_colored_heading(doc, "9.1 配置步骤", level=2)
add_step(doc, 1, "获取 API Key",
         "访问 DeepSeek 开放平台（https://platform.deepseek.com/）：\n"
         "注册账号 → 创建 API Key → 复制密钥保存。")
add_step(doc, 2, "填写配置",
         "在平台侧边栏底部找到「⚙️ API 配置」展开面板：\n"
         "① 粘贴 API Key 到输入框\n"
         "② 确认 API Base URL（默认：https://api.deepseek.com/v1）\n"
         "③ 选择模型（推荐：deepseek-chat）")
add_step(doc, 3, "保存连接",
         "点击「💾 保存配置」按钮保存设置。\n"
         "配置将加密存储在本地配置文件，下次启动自动加载。")

add_info_box(doc, "API Key 是敏感信息，请妥善保管，不要分享给他人。"
              "平台不会将您的 API Key 上传到云端。")

doc.add_page_break()

# =============================================================
# 第10章：常见问题
# =============================================================

add_colored_heading(doc, "10. 常见问题（FAQ）", level=1, color=(0, 51, 102))

faqs = [
    ("Q1：上传图片后显示红色感叹号？",
     "点击「重新选择」按钮后重新上传即可。"
     "建议使用 JPG 或 PNG 格式的图片，避免使用特殊字符文件名。"),
    ("Q2：血管分割结果是全白或全绿的？",
     "请确保上传的是标准眼底照片（非已处理过的图像）。"
     "系统会自动增强图像对比度以提高分割精度。"),
    ("Q3：AI 科研导师不回复？",
     "请检查：① API Key 是否已正确配置；"
     "② 网络是否正常（需要访问 DeepSeek API）；"
     "③ API 账户是否有可用余额。"),
    ("Q4：分析速度慢？",
     "首次运行时需要加载模型（约 5-10 秒），后续分析会更快。"
     "图像尺寸过大也会影响处理速度，建议使用 2000×2000 像素以内的图像。"),
    ("Q5：支持哪些图片格式？",
     "支持的格式：JPG/JPEG、PNG、BMP、TIFF/TIF、WebP。"
     "文件扩展名不区分大小写。"),
    ("Q6：如何更换深色/浅色主题？",
     "平台默认使用深色主题。如需切换，点击右上角 Settings 菜单，"
     "在 Theme 选项中选择 Light/Dark。"),
    ("Q7：学习分析数据存储在哪里？",
     "学习数据存储在本地 SQLite 数据库中，文件位于：\n"
     "项目目录/learning_analytics.db"),
    ("Q8：平台支持多人同时使用吗？",
     "Streamlit 默认支持多用户并发访问。"
     "每个用户的会话状态（上传的图像、分析结果）相互独立。"),
]

for question, answer in faqs:
    p_q = doc.add_paragraph()
    run_q = p_q.add_run(question)
    run_q.bold = True
    run_q.font.size = Pt(11)
    run_q.font.color.rgb = RGBColor(0, 51, 102)

    p_a = doc.add_paragraph(answer)
    p_a.paragraph_format.left_indent = Cm(0.5)
    for run in p_a.runs:
        run.font.size = Pt(10)

    doc.add_paragraph()  # spacing between FAQs

# =============================================================
# 页脚：版权信息
# =============================================================

doc.add_paragraph()
divider = doc.add_paragraph("─" * 60)
divider.alignment = WD_ALIGN_PARAGRAPH.CENTER

footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_run = footer.add_run("FundusAI-Edu  ·  眼底图像AI教学与科研辅助平台  ·  2026")
footer_run.font.size = Pt(9)
footer_run.font.color.rgb = RGBColor(150, 150, 150)

# ── 保存 ──
doc.save(OUTPUT)
print(f"[OK] User manual generated: {OUTPUT}")
