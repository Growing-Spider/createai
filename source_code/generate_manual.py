# =============================================================
# generate_manual.py — 生成 FundusAI-Edu 用户使用手册 (DOCX)
# 输出: F:/demo/source_code/FundusAI-Edu_用户使用手册.docx
# =============================================================

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from pathlib import Path

OUTPUT_PATH = Path(__file__).parent / "FundusAI-Edu_用户使用手册.docx"


def set_style(doc):
    """设置文档默认样式"""
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # 标题样式
    for i in range(1, 4):
        heading_style = doc.styles[f'Heading {i}']
        heading_font = heading_style.font
        heading_font.color.rgb = RGBColor(0x0D, 0x11, 0x17)
        heading_font.bold = True
        if i == 1:
            heading_font.size = Pt(22)
        elif i == 2:
            heading_font.size = Pt(16)
        else:
            heading_font.size = Pt(13)


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_para(doc, text, bold=False, italic=False, indent=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if indent:
        p.paragraph_format.left_indent = Cm(1)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    if level > 0:
        p.paragraph_format.left_indent = Cm(1 + level * 1.5)
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True

    # 数据行
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            table.rows[r_idx + 1].cells[c_idx].text = str(val)

    doc.add_paragraph()  # 空行
    return table


def add_tip(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(f"[提示] {text}")
    run.italic = True
    run.font.color.rgb = RGBColor(0x58, 0xA6, 0xFF)
    p.paragraph_format.left_indent = Cm(1)
    return p


def build():
    doc = Document()
    set_style(doc)

    # ============================================================
    # 封面
    # ============================================================
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("FundusAI-Edu")
    run.bold = True
    run.font.size = Pt(36)
    run.font.color.rgb = RGBColor(0x58, 0xA6, 0xFF)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("眼底图像AI教学与科研辅助平台")
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x8B, 0x94, 0x9E)

    doc.add_paragraph()

    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub2.add_run("用户使用手册 v1.0")
    run.font.size = Pt(14)

    sub3 = doc.add_paragraph()
    sub3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub3.add_run("2026年6月")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x8B, 0x94, 0x9E)

    doc.add_page_break()

    # ============================================================
    # 目录页
    # ============================================================
    add_heading(doc, "目录", level=1)
    toc_items = [
        "1. 平台简介",
        "2. 安装与启动",
        "   2.1 环境要求",
        "   2.2 安装步骤",
        "   2.3 启动方式",
        "3. 功能模块",
        "   3.1 眼底图像智能分析",
        "   3.2 血管拓扑特征分析",
        "   3.3 AI科研导师 (RAG问答)",
        "   3.4 科研选题生成器",
        "   3.5 论文写作辅助",
        "   3.6 学习分析",
        "4. 侧边栏设置",
        "5. 常见问题",
        "6. 技术支持",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(4)

    doc.add_page_break()

    # ============================================================
    # 1. 平台简介
    # ============================================================
    add_heading(doc, "1. 平台简介", level=1)
    add_para(doc, "FundusAI-Edu 是一套面向医学教学与科研的眼底图像人工智能分析平台，将深度学习、图像处理和自然语言处理技术整合为一站式工具，帮助师生快速掌握眼底图像分析方法并产出科研成果。")
    add_para(doc, "平台核心能力：")
    add_bullet(doc, "眼底图像血管自动分割与可视化")
    add_bullet(doc, "血管网络拓扑特征六维分析（密度、分叉点、终末点、平均宽度、迂曲度、分形维数）")
    add_bullet(doc, "糖尿病视网膜病变（DR）风险自动分级")
    add_bullet(doc, "基于 RAG 的医学知识智能问答")
    add_bullet(doc, "AI 辅助科研选题、研究方案生成")
    add_bullet(doc, "论文框架构建与章节写作辅助")
    add_bullet(doc, "用户学习行为统计与分析画像")

    add_para(doc, "适用对象：高职高专 / 本科生 / 硕士研究生 / 博士研究生。涵盖眼科、医学影像、生物医学工程等多个方向。")

    doc.add_page_break()

    # ============================================================
    # 2. 安装与启动
    # ============================================================
    add_heading(doc, "2. 安装与启动", level=1)

    add_heading(doc, "2.1 环境要求", level=2)
    add_table(doc, ["项目", "要求"], [
        ["操作系统", "Windows 10 / 11 (64位)"],
        ["Python", "3.10 及以上"],
        ["内存", "建议 8GB 及以上"],
        ["硬盘空间", "建议 5GB 可用空间"],
        ["GPU", "可选，CPU 模式也可正常运行"],
        ["浏览器", "Edge / Chrome / Firefox（弹窗模式自动调用 Edge WebView2）"],
    ])

    add_heading(doc, "2.2 安装步骤", level=2)
    add_para(doc, "方式一：源码运行（推荐）", bold=True)
    add_bullet(doc, "打开终端，进入项目目录")
    add_bullet(doc, "创建虚拟环境：python -m venv venv")
    add_bullet(doc, "激活环境：venv\\Scripts\\activate")
    add_bullet(doc, "安装依赖：pip install -r requirements.txt")
    add_bullet(doc, "启动应用：streamlit run app.py")
    add_para(doc, "")
    add_para(doc, "方式二：可执行文件运行", bold=True)
    add_bullet(doc, "将 FundusAI-Edu.exe 和配套文件夹放在同一目录")
    add_bullet(doc, "双击 FundusAI-Edu.exe 启动")
    add_bullet(doc, "等待约 10-20 秒，弹窗自动打开")

    add_heading(doc, "2.3 启动方式", level=2)
    add_table(doc, ["命令", "效果"], [
        ["streamlit run app.py", "浏览器网页模式，自动打开默认浏览器"],
        ["python run.py", "桌面弹窗模式，独立原生窗口"],
        ["双击 FundusAI-Edu.exe", "打包后的可执行文件，无需安装 Python"],
    ])

    add_tip(doc, "弹窗模式依赖 pywebview 和 Edge WebView2，Win10/11 自带无需额外安装。")

    doc.add_page_break()

    # ============================================================
    # 3. 功能模块
    # ============================================================
    add_heading(doc, "3. 功能模块", level=1)
    add_para(doc, "平台共包含 6 个功能模块，通过左侧导航切换。下面逐一介绍各模块的使用方法。")

    # --- 3.1 ---
    add_heading(doc, "3.1 眼底图像智能分析", level=2)
    add_para(doc, "功能说明：上传眼底图像，自动完成血管分割、病变识别和拓扑特征提取。", bold=True)
    add_para(doc, "操作步骤：")
    add_bullet(doc, "切换到「🔬 眼底图像智能分析」页面")
    add_bullet(doc, "点击「选择眼底图像」上传一张眼底照片（支持 JPG / PNG / BMP / TIFF / WebP 格式）")
    add_bullet(doc, "上传后右侧自动显示原始图像预览和文件元数据")
    add_bullet(doc, "点击「🚀 开始分析」按钮启动 AI 分析")
    add_bullet(doc, "等待进度条走完（约 5-15 秒），查看分析结果")
    add_para(doc, "结果解读：")
    add_table(doc, ["选项卡", "显示内容"], [
        ["血管分割", "左侧：血管掩码图（绿色血管）; 右侧：血管叠加图（原图+绿色血管标注）; 下方：骨架可视化三连图"],
        ["病变识别", "DR 分级（No DR / Mild / Moderate / Severe DR）、置信度、出血风险、渗出风险"],
        ["拓扑特征", "六大指标：血管密度、分叉点数、终末点数、平均血管宽度、迂曲度、分形维数，以及雷达图"],
    ])

    add_tip(doc, "无 U-Net 预训练权重时，平台自动切换 Frangi 滤波器演示模式，功能完全可用。")
    add_tip(doc, "分析前可使用「🗑️ 清除」按钮重置已上传的图片，重新选择。")

    # --- 3.2 ---
    add_heading(doc, "3.2 血管拓扑特征分析", level=2)
    add_para(doc, "功能说明：深度解析视网膜血管网络的拓扑结构，自动计算六大特征并生成可视化图表。", bold=True)
    add_para(doc, "使用前提：先在「眼底图像智能分析」中完成一次分析。分析结果会跨页面保留。")
    add_para(doc, "特征解读参考：")
    add_table(doc, ["特征", "单位", "临床意义"], [
        ["血管密度", "比例 (0-1)", "反映视网膜血管覆盖率，异常低值可能提示缺血"],
        ["分叉点数量", "个", "血管分支交点数量，反映血管网络复杂度"],
        ["终末点数量", "个", "血管末梢数量，与血管生长有关"],
        ["平均血管宽度", "像素", "CAD 值：正常约 5-8px，偏宽提示血管病变"],
        ["迂曲度", "比值 (>1)", "血管弯曲程度，>2.0 提示异常迂曲"],
        ["分形维数", "(1-2)", "血管网络空间填充效率，"],
    ])

    add_tip(doc, "骨架可视化三连图：原始图、血管分割（绿色）、骨架叠加（红色），直观对比血管网络结构。")

    # --- 3.3 ---
    add_heading(doc, "3.3 AI 科研导师（RAG 问答）", level=2)
    add_para(doc, "功能说明：基于内置眼底医学知识库，结合 DeepSeek 大模型提供专业智能问答。", bold=True)
    add_para(doc, "使用方式：")
    add_bullet(doc, "在对话框中输入医学/科研相关问题")
    add_bullet(doc, "系统自动检索知识库，结合大模型智能回答")
    add_bullet(doc, "支持多轮对话，历史记录保留在右侧")
    add_para(doc, "知识库管理：")
    add_bullet(doc, "在侧边栏「📚 知识库管理」上传 PDF/TXT/DOCX 文档扩充知识库")
    add_bullet(doc, "点击「🔄 重建知识库」使新文档生效")
    add_tip(doc, "未配置 DeepSeek API Key 时，使用本地关键词搜索模式，功能降级但不报错。")

    # --- 3.4 ---
    add_heading(doc, "3.4 科研选题生成器", level=2)
    add_para(doc, "功能说明：根据研究方向、学术层次和自定义需求，AI 自动生成个性化科研选题和实验方案。", bold=True)
    add_para(doc, "操作步骤：")
    add_bullet(doc, "选择「学术层次」：高职高专 / 本科生 / 硕士研究生 / 博士研究生")
    add_bullet(doc, "输入「研究方向」（如：糖尿病视网膜病变的AI诊断）")
    add_bullet(doc, "设置「生成选题数量」（1-5个）")
    add_bullet(doc, "可选填「其他要求」（如：希望有公开数据集可用）")
    add_bullet(doc, "点击「生成选题」获取 AI 推荐的课题列表")
    add_bullet(doc, "选择一个选题后，可继续「生成研究方案」和「设计实验报告」")
    add_tip(doc, "学术层次决定选题难度和深度。高职高专偏重应用型，博士生偏重理论创新。")

    # --- 3.5 ---
    add_heading(doc, "3.5 论文写作辅助", level=2)
    add_para(doc, "功能说明：AI 辅助论文框架搭建、各章节内容生成和润色优化。", bold=True)
    add_para(doc, "支持的章节类型：")
    add_bullet(doc, "摘要 (Abstract)")
    add_bullet(doc, "引言 (Introduction)")
    add_bullet(doc, "相关工作 (Related Work)")
    add_bullet(doc, "方法 (Method)")
    add_bullet(doc, "实验与结果 (Experiments & Results)")
    add_bullet(doc, "讨论 (Discussion)")
    add_bullet(doc, "结论 (Conclusion)")
    add_bullet(doc, "参考文献推荐")
    add_bullet(doc, "期刊推荐")
    add_tip(doc, "建议先使用「科研选题生成器」确定课题方向，再利用本模块辅助论文写作，效率更高。")

    # --- 3.6 ---
    add_heading(doc, "3.6 学习分析", level=2)
    add_para(doc, "功能说明：自动记录用户平台操作行为，生成可视化学习分析报告。", bold=True)
    add_para(doc, "统计维度：")
    add_bullet(doc, "总操作次数与分析次数")
    add_bullet(doc, "各模块使用频率分布")
    add_bullet(doc, "分析图像数量和类型统计")
    add_bullet(doc, "学习时长趋势")
    add_bullet(doc, "个人学习画像（进步曲线）")

    doc.add_page_break()

    # ============================================================
    # 4. 侧边栏设置
    # ============================================================
    add_heading(doc, "4. 侧边栏设置", level=1)
    add_table(doc, ["设置项", "说明"], [
        ["用户设置", "可自定义用户名，用于学习记录"],
        ["API 配置", "输入 DeepSeek API Key 以启用 AI 功能（访问 platform.deepseek.com 获取）"],
        ["知识库管理", "上传 PDF/TXT/DOCX 扩充知识库; 点击重建使新文档生效"],
        ["快速统计", "显示当前用户的累计操作和分析次数"],
    ])

    add_tip(doc, "不配置 API Key 的情况下，AI 功能使用本地离线预设内容，不影响页面交互体验。")

    doc.add_page_break()

    # ============================================================
    # 5. 常见问题
    # ============================================================
    add_heading(doc, "5. 常见问题", level=1)

    faq = [
        ("Q: 为什么不显示血管分割结果？",
         "A: 如果掩码图全黑或全白，可能是图像对比度异常。请尝试：1) 检查图像是否为标准眼底照片 2) 调整图像亮度/对比度后重试 3) 如仍不行，联系技术支持。"),
        ("Q: Connection lost 错误怎么办？",
         "A: 此为 WebSocket 连接断开。关闭页面/窗口重新打开即可。如频繁出现，检查防火墙是否阻止了本地 8501 端口。"),
        ("Q: 可以不用 GPU 运行吗？",
         "A: 可以。平台默认使用 CPU 模式，Frangi 滤波器在没有 U-Net 权重时也能正常工作。"),
        ("Q: 支持哪些图像格式？",
         "A: JPG、JPEG、PNG、BMP、TIFF、TIF、WebP。不区分大小写。"),
        ("Q: DeepSeek API Key 如何获取？",
         "A: 访问 https://platform.deepseek.com 注册账号，在「API Keys」页面创建密钥。"),
        ("Q: 知识库文档有什么要求？",
         "A: 支持 PDF、TXT、DOCX 格式，建议文档内容与眼底医学相关。上传后在「知识库管理」中点击重建生效。"),
    ]
    for q, a in faq:
        add_para(doc, q, bold=True)
        add_para(doc, a, indent=True)

    doc.add_page_break()

    # ============================================================
    # 6. 技术支持
    # ============================================================
    add_heading(doc, "6. 技术支持", level=1)
    add_para(doc, "如遇技术问题，请通过以下方式联系我们：")
    add_bullet(doc, "项目仓库：查看源码中的文档和更新日志")
    add_bullet(doc, "技术支持邮箱：developer@fundus-ai.edu")
    add_bullet(doc, "DeepSeek API 平台：https://platform.deepseek.com")
    add_para(doc, "")
    add_para(doc, "版本：v1.0 (弹窗版本)")
    add_para(doc, "更新日期：2026年6月")
    add_para(doc, "© 2026 FundusAI-Edu 团队")

    # ---- 保存 ----
    doc.save(str(OUTPUT_PATH))
    print(f"[OK] 使用手册已生成: {OUTPUT_PATH}")


if __name__ == "__main__":
    build()
