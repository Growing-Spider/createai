# =============================================================
# generate_install_manual.py — FundusAI-Edu 安装部署手册 (DOCX)
# 输出: F:/demo/source_code/FundusAI-Edu_安装部署手册.docx
# =============================================================

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from pathlib import Path

OUTPUT_PATH = Path(__file__).parent / "FundusAI-Edu_安装部署手册.docx"


def set_style(doc):
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Consolas'
    font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    for i in range(1, 4):
        hs = doc.styles[f'Heading {i}']
        hs.font.color.rgb = RGBColor(0x0D, 0x11, 0x17)
        hs.font.bold = True
        hs.font.size = {1: Pt(22), 2: Pt(16), 3: Pt(13)}[i]

    # 代码块样式
    code_style = doc.styles.add_style('CodeBlock', 1)
    code_style.font.name = 'Consolas'
    code_style.font.size = Pt(9)
    code_style.font.color.rgb = RGBColor(0x58, 0xA6, 0xFF)
    code_style.paragraph_format.space_before = Pt(4)
    code_style.paragraph_format.space_after = Pt(4)
    code_style.paragraph_format.left_indent = Cm(1)


def add_h(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_p(doc, text, bold=False, indent=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    if indent:
        p.paragraph_format.left_indent = Cm(1)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    if level > 0:
        p.paragraph_format.left_indent = Cm(1 + level * 1.5)


def add_code(doc, code):
    """添加代码块"""
    for line in code.strip().split('\n'):
        p = doc.add_paragraph(line, style='CodeBlock')


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            table.rows[r_idx + 1].cells[c_idx].text = str(val)
    doc.add_paragraph()
    return table


def add_note(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(f"[注意] {text}")
    run.italic = True
    run.font.color.rgb = RGBColor(0xF7, 0x81, 0x66)
    p.paragraph_format.left_indent = Cm(1)


def add_tip(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(f"[提示] {text}")
    run.italic = True
    run.font.color.rgb = RGBColor(0x58, 0xA6, 0xFF)
    p.paragraph_format.left_indent = Cm(1)


def build():
    doc = Document()
    set_style(doc)

    # ============ 封面 ============
    for _ in range(6):
        doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("FundusAI-Edu")
    r.bold = True; r.font.size = Pt(36); r.font.color.rgb = RGBColor(0x58, 0xA6, 0xFF)

    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = s.add_run("安装部署手册")
    r.font.size = Pt(20); r.font.color.rgb = RGBColor(0x8B, 0x94, 0x9E)

    doc.add_paragraph()
    for line in ["版本 v1.0", "2026年6月", "适用环境: Windows 10/11"]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.font.size = Pt(12); run.font.color.rgb = RGBColor(0x8B, 0x94, 0x9E)

    doc.add_page_break()

    # ============ 目录 ============
    add_h(doc, "目录")
    for item in [
        "1. 系统要求",
        "2. 快速安装（5分钟）",
        "   2.1 安装 Python",
        "   2.2 创建虚拟环境",
        "   2.3 安装依赖包",
        "3. 启动应用",
        "   3.1 浏览器模式",
        "   3.2 桌面弹窗模式",
        "4. 打包为 EXE",
        "   4.1 打包步骤",
        "   4.2 分发包结构",
        "5. 依赖包清单",
        "6. 常见安装问题",
    ]:
        p = doc.add_paragraph(item); p.paragraph_format.space_after = Pt(3)
    doc.add_page_break()

    # ============ 1. 系统要求 ============
    add_h(doc, "1. 系统要求")

    add_table(doc, ["项目", "最低要求", "推荐配置"], [
        ["操作系统", "Windows 10 (64位)", "Windows 11 (64位)"],
        ["Python", "3.10", "3.13"],
        ["内存", "4 GB", "8 GB 及以上"],
        ["硬盘空间", "3 GB", "5 GB（含模型权重）"],
        ["GPU", "不需要", "NVIDIA GPU（可选加速）"],
        ["浏览器", "Edge / Chrome", "Edge（弹窗模式必需）"],
        ["网络", "不需要（离线可用）", "需联网（AI 功能）"],
    ])

    add_note(doc, "如仅使用演示模式（Frangi 滤波器），无需 GPU 且无需下载预训练模型权重。")

    doc.add_page_break()

    # ============ 2. 快速安装 ============
    add_h(doc, "2. 快速安装（约 5 分钟）")

    add_h(doc, "2.1 安装 Python", 2)
    add_p(doc, "方式一：官网下载（推荐）")
    add_bullet(doc, "访问 https://www.python.org/downloads/")
    add_bullet(doc, "下载 Python 3.10 或更高版本（Windows installer 64-bit）")
    add_bullet(doc, "安装时勾选「Add Python to PATH」")
    add_p(doc, "")
    add_p(doc, "方式二：使用 Anaconda（如已安装）")
    add_code(doc, "conda create -n fundus-ai python=3.10 -y\nconda activate fundus-ai")
    add_p(doc, "")
    add_p(doc, "验证安装：")
    add_code(doc, "python --version\npip --version")

    add_h(doc, "2.2 创建虚拟环境", 2)
    add_p(doc, "在项目根目录下执行：")
    add_code(doc, "cd F:\\demo\\source_code\npython -m venv venv")
    add_p(doc, "激活虚拟环境：")
    add_code(doc, "# Windows CMD:\nvenv\\Scripts\\activate\n\n# Windows PowerShell:\n.\\venv\\Scripts\\Activate.ps1\n\n# Git Bash:\nsource venv/Scripts/activate")
    add_tip(doc, "终端提示符前出现 (venv) 表示激活成功。")

    add_h(doc, "2.3 安装依赖包", 2)
    add_p(doc, "方法一：一键安装（推荐）")
    add_code(doc, "pip install -r requirements.txt")
    add_p(doc, "方法二：分步安装（如某包安装失败）")
    add_p(doc, "先装 PyTorch CPU 版：", bold=True)
    add_code(doc, "pip install torch torchvision --index-url https://download.pytorch.org/whl/cpu")
    add_p(doc, "再装核心依赖：", bold=True)
    add_code(doc, "pip install streamlit opencv-python scikit-image scipy\npip install pandas matplotlib seaborn plotly\npip install langchain langchain-community langchain-openai\npip install chromadb sentence-transformers\npip install openai networkx Pillow pypdf python-docx reportlab\npip install python-dotenv tqdm loguru requests")
    add_p(doc, "弹窗模式额外依赖：", bold=True)
    add_code(doc, "pip install pywebview")
    add_tip(doc, "如网络慢，加国内镜像：pip install xxx -i https://pypi.tuna.tsinghua.edu.cn/simple")

    doc.add_page_break()

    # ============ 3. 启动应用 ============
    add_h(doc, "3. 启动应用")

    add_h(doc, "3.1 浏览器模式", 2)
    add_p(doc, "在终端中执行：")
    add_code(doc, "cd F:\\demo\\source_code\nstreamlit run app.py")
    add_p(doc, "浏览器自动打开 http://localhost:8501，即可使用。")
    add_p(doc, "如需局域网访问：")
    add_code(doc, "streamlit run app.py --server.address 0.0.0.0")

    add_h(doc, "3.2 桌面弹窗模式", 2)
    add_p(doc, "前提：已安装 pywebview，系统自带 Edge WebView2。")
    add_code(doc, "cd F:\\demo\\source_code\npython run.py")
    add_p(doc, "弹出一个独立原生窗口，关闭窗口后服务自动停止。")
    add_p(doc, "也可用专用入口（PyInstaller 兼容）：")
    add_code(doc, "python launcher.py")

    add_table(doc, ["启动方式", "命令", "界面"], [
        ["浏览器", "streamlit run app.py", "浏览器标签页"],
        ["弹窗", "python run.py", "独立桌面窗口"],
        ["EXE", "双击 FundusAI-Edu.exe", "独立桌面窗口"],
    ])

    doc.add_page_break()

    # ============ 4. 打包为 EXE ============
    add_h(doc, "4. 打包为可执行文件 (EXE)")

    add_h(doc, "4.1 打包步骤", 2)
    add_p(doc, "确保已完成第2章的安装，然后：")
    add_p(doc, "步骤 1：安装 PyInstaller", bold=True)
    add_code(doc, "pip install pyinstaller")
    add_p(doc, "步骤 2：执行打包", bold=True)
    add_code(doc, "# 方式一：一键脚本\npython build.py\n\n# 方式二：直接使用 spec 文件\npyinstaller FundusAI-Edu.spec")
    add_p(doc, "步骤 3：等待完成", bold=True)
    add_p(doc, "打包耗时约 10-30 分钟（取决于机器性能）。成功后输出：")
    add_code(doc, "F:\\demo\\dist\\FundusAI-Edu.exe")
    add_tip(doc, "EXE 体积约 500MB-1.5GB，主要来自 PyTorch 和 OpenCV 库。")

    add_h(doc, "4.2 分发包结构", 2)
    add_p(doc, "打包后需将以下文件放在同一目录：")
    add_code(doc, "FundusAI-Edu.exe       # 主程序\n├── app.py              # Streamlit 入口\n├── config.py           # 配置文件\n├── modules/            # 功能模块 (6个)\n├── knowledge_base/     # 知识库文件\n├── models/             # 模型目录\n└── .streamlit/         # Streamlit 配置")
    add_note(doc, "所有配套文件由打包脚本自动放置到 dist 目录，无需手动复制。")

    doc.add_page_break()

    # ============ 5. 依赖包清单 ============
    add_h(doc, "5. 依赖包清单")

    add_table(doc, ["分类", "包名", "版本", "用途"], [
        ["框架", "streamlit", ">=1.32", "Web UI 框架"],
        ["深度学习", "torch", ">=2.0", "U-Net 模型推理"],
        ["深度学习", "torchvision", ">=0.15", "图像预处理工具"],
        ["图像处理", "opencv-python", ">=4.8", "图像读写和预处理"],
        ["图像处理", "scikit-image", ">=0.22", "血管分割算法"],
        ["图像处理", "scipy", ">=1.11", "科学计算"],
        ["图像处理", "Pillow", ">=10.0", "Python 图像库"],
        ["数据", "numpy", ">=1.24", "数组运算"],
        ["数据", "pandas", ">=2.0", "数据分析"],
        ["可视化", "matplotlib", ">=3.7", "图表生成"],
        ["可视化", "seaborn", ">=0.12", "统计图表"],
        ["可视化", "plotly", ">=5.15", "交互式图表"],
        ["图分析", "networkx", ">=3.1", "血管拓扑图"],
        ["AI/LLM", "openai", ">=1.14", "DeepSeek API"],
        ["AI/LLM", "langchain", ">=0.1", "LLM 框架"],
        ["AI/LLM", "langchain-community", ">=0.0.32", "社区扩展"],
        ["AI/LLM", "langchain-openai", ">=0.1", "OpenAI 集成"],
        ["AI/LLM", "chromadb", ">=0.4", "向量数据库"],
        ["AI/LLM", "sentence-transformers", ">=2.5", "文本嵌入"],
        ["文档", "pypdf", ">=4.0", "PDF 解析"],
        ["文档", "python-docx", ">=1.0", "DOCX 处理"],
        ["文档", "reportlab", ">=4.0", "PDF 生成"],
        ["桌面", "pywebview", ">=5.0", "弹窗 GUI"],
        ["工具", "python-dotenv", ">=1.0", "环境变量"],
        ["工具", "tqdm", ">=4.66", "进度条"],
        ["工具", "loguru", ">=0.7", "日志"],
        ["工具", "requests", ">=2.31", "HTTP 请求"],
        ["打包", "pyinstaller", ">=6.0", "EXE 打包"],
    ])

    doc.add_page_break()

    # ============ 6. 常见安装问题 ============
    add_h(doc, "6. 常见安装问题")

    faq = [
        ("Q: pip install 时提示 Microsoft Visual C++ 14.0 is required",
         "A: 下载安装 Visual Studio Build Tools:\n"
         "   https://visualstudio.microsoft.com/visual-cpp-build-tools/\n"
         "   安装时勾选「C++ 生成工具」即可。"),
        ("Q: torch 安装失败或速度太慢",
         "A: 使用国内镜像：\n"
         "   pip install torch torchvision --index-url https://download.pytorch.org/whl/cpu\n"
         "   如仍失败，可到 https://pytorch.org 手动下载 wheel 文件安装。"),
        ("Q: streamlit 启动后页面空白",
         "A: 检查端口是否被占用：netstat -ano | findstr 8501\n"
         "   如有其他程序占用，可更换端口：streamlit run app.py --server.port 8502"),
        ("Q: pywebview 弹窗不显示",
         "A: 1) 确保已安装 pywebview  \n"
         "   2) 确保系统已安装 Edge WebView2（Win10/11 自带）\n"
         "   3) 如仍不行，改用浏览器模式：streamlit run app.py"),
        ("Q: PyInstaller 打包后运行报错 GBK 编码",
         "A: 项目代码已做编码兼容处理。如仍遇到，在启动命令前加：\n"
         "   set PYTHONIOENCODING=utf-8 && FundusAI-Edu.exe"),
        ("Q: 提示 ModuleNotFoundError: No module named 'xxx'",
         "A: 确认虚拟环境已激活（提示符前有 (venv)），然后：\n"
         "   pip install -r requirements.txt"),
        ("Q: chromadb 安装失败",
         "A: chromadb 需要 C++ 编译环境。可尝试：\n"
         "   pip install chromadb --no-build-isolation\n"
         "   如仅需基础功能，可跳过 chromadb，RAG 自动降级为关键词搜索。"),
    ]
    for q, a in faq:
        add_p(doc, q, bold=True)
        add_p(doc, a, indent=True)
        doc.add_paragraph()

    # ============ 尾页 ============
    doc.add_page_break()
    add_h(doc, "技术支持")
    add_bullet(doc, "文档随源码更新，最新版见项目仓库")
    add_bullet(doc, "技术邮箱：developer@fundus-ai.edu")
    add_p(doc, "")
    add_p(doc, "© 2026 FundusAI-Edu | 医学AI教育平台")

    # 保存
    doc.save(str(OUTPUT_PATH))
    print(f"[OK] 安装手册已生成: {OUTPUT_PATH}")


if __name__ == "__main__":
    build()
